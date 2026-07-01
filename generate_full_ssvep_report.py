from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from scipy import stats


ROOT = Path(".")
RESULT_DIR = ROOT / "ssvep_results" / "full_benchmark"
DOCX_PATH = ROOT / "基于优化TDCA算法的SSVEP脑电识别技术创新报告.docx"


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 13, bold=True)
    return p


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    set_run_font(r, 9, bold=bold)


def add_df_table(doc: Document, df: pd.DataFrame, title: str | None = None):
    if title:
        add_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        set_cell_text(table.rows[0].cells[i], col, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.3f}"
            else:
                text = str(val)
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def add_page_break(doc: Document):
    doc.add_page_break()


def method_summary_table(method_df: pd.DataFrame) -> pd.DataFrame:
    df = method_df.copy()
    df["准确率(%)"] = df["mean_accuracy"] * 100
    df["标准差(%)"] = df["std_accuracy"] * 100
    df["ITR(bits/min)"] = df["mean_itr_bits_min"]
    df["Top-2间隔"] = df["mean_margin"]
    return df[["window_sec", "method", "准确率(%)", "标准差(%)", "ITR(bits/min)", "Top-2间隔", "subjects"]].rename(
        columns={"window_sec": "窗口(s)", "method": "算法", "subjects": "被试数"}
    )


def best_table(method_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for window, sub in method_df.groupby("window_sec"):
        best_acc = sub.loc[sub["mean_accuracy"].idxmax()]
        best_itr = sub.loc[sub["mean_itr_bits_min"].idxmax()]
        rows.append(
            {
                "窗口(s)": window,
                "最高准确率算法": best_acc["method"],
                "最高准确率(%)": best_acc["mean_accuracy"] * 100,
                "最高ITR算法": best_itr["method"],
                "最高ITR(bits/min)": best_itr["mean_itr_bits_min"],
            }
        )
    return pd.DataFrame(rows)


def paired_stats(subject_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for window, sub in subject_df.groupby("window_sec"):
        pivot = sub.pivot(index="subject", columns="method", values="accuracy")
        for method in ["ImprovedTDCA", "TriBranchTDCA"]:
            if method not in pivot.columns or "TDCA" not in pivot.columns:
                continue
            diff = pivot[method] - pivot["TDCA"]
            try:
                t_p = stats.ttest_rel(pivot[method], pivot["TDCA"], nan_policy="omit").pvalue
            except Exception:
                t_p = np.nan
            try:
                w_p = stats.wilcoxon(diff).pvalue if np.any(np.abs(diff) > 1e-12) else 1.0
            except Exception:
                w_p = np.nan
            rows.append(
                {
                    "窗口(s)": window,
                    "比较": f"{method} - TDCA",
                    "平均提升(百分点)": float(diff.mean() * 100),
                    "提升被试数": int((diff > 0).sum()),
                    "持平被试数": int((np.abs(diff) <= 1e-12).sum()),
                    "下降被试数": int((diff < 0).sum()),
                    "配对t检验p": float(t_p),
                    "Wilcoxon p": float(w_p),
                }
            )
    return pd.DataFrame(rows)


def add_images(doc: Document):
    images = [
        ("完整 Benchmark 平均准确率", RESULT_DIR / "full_accuracy.png"),
        ("完整 Benchmark 平均 ITR", RESULT_DIR / "full_itr.png"),
        ("完整 Benchmark Top-2 分类间隔", RESULT_DIR / "full_margin.png"),
    ]
    for title, path in images:
        if path.exists():
            add_heading(doc, title, 2)
            doc.add_picture(str(path), width=Inches(6.3))
            doc.add_paragraph()


def delete_old_reports(final_path: Path):
    patterns = [
        "report_ssvep_tdca_results*.md",
        "report_ssvep_tdca_results*.txt",
        "report_ssvep_tdca_results*.docx",
        "SSVEP_TDCA_result_report*.md",
        "SSVEP_TDCA_result_report*.txt",
        "FINAL_SSVEP_TDCA*.docx",
    ]
    for pattern in patterns:
        for path in ROOT.glob(pattern):
            if path.resolve() != final_path.resolve() and path.exists():
                path.unlink()


def main():
    method_df = pd.read_csv(RESULT_DIR / "method_summary.csv")
    subject_df = pd.read_csv(RESULT_DIR / "subject_summary.csv")
    block_df = pd.read_csv(RESULT_DIR / "block_metrics.csv")

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于优化 TDCA 算法的 SSVEP 脑电识别技术创新报告")
    set_run_font(r, 18, bold=True)

    add_heading(doc, "一、实验数据集", 1)
    add_paragraph(
        doc,
        "本实验使用本地 Benchmark SSVEP 数据集，路径为 E:/世界机器人大赛/benchmark。"
        "数据共包含 35 个被试，每个被试文件的 data 字段维度为 64 × 1500 × 40 × 6，"
        "分别对应 64 个 EEG 通道、1500 个采样点、40 个刺激目标和 6 个实验 block。采样率为 250 Hz。"
    )
    add_paragraph(
        doc,
        "40 个目标频率按 Benchmark 数据维度中的实际顺序排列："
        "8, 9, ..., 15, 8.2, 9.2, ..., 15.2, ..., 8.8, 9.8, ..., 15.8 Hz。"
        "本报告使用全部 35 个被试完成完整验证。"
    )

    add_heading(doc, "二、预处理与验证方式", 1)
    add_paragraph(
        doc,
        "实验采用 9 个枕顶区 SSVEP 常用通道：PZ、PO5、PO3、POZ、PO4、PO6、O1、OZ、O2。"
        "切窗时跳过 0.5 s 预刺激段，并补偿 0.14 s 视觉延迟，然后分别截取 0.3 s、0.5 s、1 s、2 s、4 s 的 EEG 片段。"
    )
    add_paragraph(
        doc,
        "验证采用被试内留一 block 交叉验证。每次选择 1 个 block 作为测试集，其余 5 个 block 作为训练集。"
        "每个被试每个窗口每个算法测试 240 个 trial；完整实验共得到 5250 条 block 级结果。"
    )

    add_heading(doc, "三、对比算法", 1)
    methods = [
        "CCA：标准典型相关分析，使用 5 个谐波构造正余弦参考信号。",
        "FBCCA：滤波器组 CCA，使用多个子带滤波后进行相关分数融合。",
        "TDCA：论文式 TDCA，包括滤波器组、时间延迟嵌入、参考信号投影增强、DSP 判别空间滤波和模板相关分类。",
        "ImprovedTDCA：TDCA 判别模板分支 + FBCCA 参考相关分支 + 训练集内融合权重校准 + 2 个判别分量。",
        "TriBranchTDCA：TDCA 判别模板分支 + FBCCA 参考相关分支 + FBTRCA 跨 trial 相关分支 + 训练集内融合权重校准。",
    ]
    for item in methods:
        add_paragraph(doc, item, style="List Bullet")

    add_heading(doc, "四、完整 Benchmark 结果", 1)
    add_df_table(doc, best_table(method_df), "表 1：每个时间窗的最优算法")
    add_df_table(doc, method_summary_table(method_df), "表 2：35 个被试的完整平均结果")

    add_heading(doc, "五、相对 TDCA 的提升统计", 1)
    stats_df = paired_stats(subject_df)
    add_df_table(doc, stats_df, "表 3：ImprovedTDCA 和 TriBranchTDCA 相对标准 TDCA 的被试内配对统计")

    add_heading(doc, "六、主要结论", 1)
    add_paragraph(
        doc,
        "完整 35 被试结果显示，0.3 s 与 0.5 s 极短窗下 ImprovedTDCA 的平均准确率最高；"
        "1 s 和 2 s 下 TriBranchTDCA 的平均准确率最高；4 s 长窗下 ImprovedTDCA 略高。"
    )
    add_paragraph(
        doc,
        "在 1 s 窗口中，TDCA、ImprovedTDCA 和 TriBranchTDCA 均接近天花板，"
        "但 TriBranchTDCA 将平均准确率从 TDCA 的 95.67% 提升到 96.15%，"
        "平均 ITR 从 294.97 bits/min 提升到 297.77 bits/min，同时 Top-2 分类间隔也更大。"
    )
    add_paragraph(
        doc,
        "在 0.5 s 窗口中，ImprovedTDCA 将平均准确率从 TDCA 的 82.06% 提升到 83.30%，"
        "平均 ITR 从 460.60 bits/min 提升到 471.74 bits/min。"
        "这说明短窗实时解码中，加入 FBCCA 参考相关分支和训练集内校准融合可以带来稳定增益。"
    )
    add_paragraph(
        doc,
        "综合准确率、ITR 与短窗稳定性，本系统建议采用分段策略："
        "0.3 s、0.5 s 和 4 s 使用 ImprovedTDCA，1 s 和 2 s 使用 TriBranchTDCA。"
        "如果工程部署需要统一模型，则优先选择 TriBranchTDCA，因为其在 1 s 和 2 s 实时控制窗口上表现最佳。"
    )

    add_heading(doc, "七、结果图", 1)
    add_images(doc)

    add_heading(doc, "八、结果文件", 1)
    files = [
        "ssvep_results/full_benchmark/block_metrics.csv：block 级完整结果。",
        "ssvep_results/full_benchmark/trial_predictions.csv：trial 级预测结果。",
        "ssvep_results/full_benchmark/subject_summary.csv：每个被试的平均结果。",
        "ssvep_results/full_benchmark/method_summary.csv：35 个被试的算法汇总。",
        "ssvep_results/full_benchmark/full_accuracy.png：完整准确率图。",
        "ssvep_results/full_benchmark/full_itr.png：完整 ITR 图。",
        "ssvep_results/full_benchmark/full_margin.png：完整 Top-2 间隔图。",
    ]
    for item in files:
        add_paragraph(doc, item, style="List Bullet")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    delete_old_reports(DOCX_PATH)
    print(DOCX_PATH.resolve())


if __name__ == "__main__":
    main()
